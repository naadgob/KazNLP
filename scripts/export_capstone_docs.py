"""Export capstone markdown → .docx and WBS.csv → .xlsx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font

ROOT = Path(__file__).resolve().parents[1]
CAPSTONE = ROOT / "docs" / "capstone"


def md_to_docx(md_path: Path, docx_path: Path, title: str | None = None) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.strip() == "---":
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            continue
        if line.startswith("|") and "|" in line[1:]:
            p = doc.add_paragraph(line)
            p.runs[0].font.name = "Consolas"
            p.runs[0].font.size = Pt(9)
            continue
        if line.startswith("```"):
            continue
        # strip simple markdown bold/code
        clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
        clean = re.sub(r"`([^`]+)`", r"\1", clean)
        if clean.startswith("- "):
            doc.add_paragraph(clean[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", clean):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", clean), style="List Number")
        else:
            doc.add_paragraph(clean)

    if title:
        doc.core_properties.title = title
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Wrote {docx_path.relative_to(ROOT)}")


def csv_to_xlsx(csv_path: Path, xlsx_path: Path) -> None:
    import csv

    wb = Workbook()
    ws = wb.active
    ws.title = "WBS"
    with csv_path.open(encoding="utf-8", newline="") as f:
        for r, row in enumerate(csv.reader(f), start=1):
            for c, val in enumerate(row, start=1):
                cell = ws.cell(row=r, column=c, value=val)
                if r == 1:
                    cell.font = Font(bold=True)
    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 6
    ws.column_dimensions["C"].width = 58
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 12
    ws.column_dimensions["F"].width = 42
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(xlsx_path)
    print(f"Wrote {xlsx_path.relative_to(ROOT)}")


def main() -> None:
    md_to_docx(
        CAPSTONE / "Final_Report.md",
        CAPSTONE / "Final_Report.docx",
        "KazNLP Final Report",
    )
    md_to_docx(
        CAPSTONE / "Action_Plan.md",
        CAPSTONE / "Action_Plan.docx",
        "KazNLP Action Plan",
    )
    csv_to_xlsx(CAPSTONE / "WBS.csv", CAPSTONE / "WBS.xlsx")


if __name__ == "__main__":
    main()
